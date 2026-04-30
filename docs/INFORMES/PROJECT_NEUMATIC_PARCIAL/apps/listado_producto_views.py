# neumatic\apps\maestros\views\producto_views.py
from django.urls import reverse_lazy
from ..views.cruds_views_generics import *
from ..models.producto_models import Producto
from ..forms.producto_forms import ProductoForm

#Modulos para generar el informe

from django.http import HttpResponse
from django.views import View
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.pdfgen import canvas
from io import BytesIO, StringIO
import csv
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from django.core.mail import EmailMessage
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls




class ConfigViews():
	# Modelo
	model = Producto
	
	# Formulario asociado al modelo
	form_class = ProductoForm
	
	# Aplicación asociada al modelo
	app_label = model._meta.app_label
	
	#-- Deshabilitado por redundancia:
	# # Título del listado del modelo
	# master_title = model._meta.verbose_name_plural
	
	#-- Usar esta forma cuando el modelo esté compuesto de una sola palabra: Ej. Color.
	model_string = model.__name__.lower()  #-- Usar esta forma cuando el modelo esté compuesto de una sola palabra: Ej. Color.
	
	#-- Usar esta forma cuando el modelo esté compuesto por más de una palabra: Ej. TipoCambio colocar "tipo_cambio".
	#model_string = "tipo_cambio"
	
	# Permisos
	permission_add = f"{app_label}.add_{model_string}"
	permission_change = f"{app_label}.change_{model_string}"
	permission_delete = f"{app_label}.delete_{model_string}"
	
	# Vistas del CRUD del modelo
	list_view_name = f"{model_string}_list"
	create_view_name = f"{model_string}_create"
	update_view_name = f"{model_string}_update"
	delete_view_name = f"{model_string}_delete"
	
	# Plantilla para crear o actualizar el modelo
	template_form = f"{app_label}/{model_string}_form.html"
	
	# Plantilla para confirmar eliminación de un registro
	template_delete = "base_confirm_delete.html"
	
	# Plantilla de la lista del CRUD
	template_list = f'{app_label}/maestro_list.html'
	
	# Contexto de los datos de la lista
	context_object_name	= 'objetos'
	
	# Vista del home del proyecto
	home_view_name = "home"
	
	# Nombre de la url 
	success_url = reverse_lazy(list_view_name)


class DataViewList():
	search_fields = ['nombre_producto']
	ordering = ['nombre_producto']

	paginate_by = 8

	table_headers = {
		'codigo_producto': (2, 'Código'),
		'nombre_producto': (2, 'Nombre'),
        'tipo_producto': (2, 'Tip. Producto'),
        'precio': (2, 'Precio'),
        'costo': (1, 'Costo'),
        'stock': (1, 'Sctok'),
		'acciones': (2, 'Acciones'),
	}

	table_data = [
		{'field_name': 'codigo_producto', 'date_format': None},
        {'field_name': 'nombre_producto', 'date_format': None},
        {'field_name': 'tipo_producto', 'date_format': None},
        {'field_name': 'precio', 'date_format': None},
        {'field_name': 'costo', 'date_format': None},
        {'field_name': 'stock', 'date_format': None},
        
		
	]

# VendedorListView - Inicio
class ListadoProductoListView(MaestroListView):
	model = ConfigViews.model
	template_name = ConfigViews.template_list
	context_object_name = ConfigViews.context_object_name
	
	search_fields = DataViewList.search_fields
	ordering = DataViewList.ordering
	
	extra_context = {
		"master_title": ConfigViews.model._meta.verbose_name_plural,
		"home_view_name": ConfigViews.home_view_name,
		"list_view_name": ConfigViews.list_view_name,
		"create_view_name": ConfigViews.create_view_name,
		"update_view_name": ConfigViews.update_view_name,
		"delete_view_name": ConfigViews.delete_view_name,
		"table_headers": DataViewList.table_headers,
		"table_data": DataViewList.table_data,
	}


# ProductoCreateView - Inicio
class ListadoProductoCreateView(MaestroCreateView):
	model = ConfigViews.model
	list_view_name = ConfigViews.list_view_name
	form_class = ConfigViews.form_class
	template_name = ConfigViews.template_form
	success_url = ConfigViews.success_url
	
	#-- Indicar el permiso que requiere para ejecutar la acción.
	# (revisar de donde lo copiaste que tienes asignado permission_change en vez de permission_add)
	permission_required = ConfigViews.permission_add
	
	extra_context = {
		"accion": f"Editar {ConfigViews.model._meta.verbose_name}",
		"list_view_name" : ConfigViews.list_view_name
	}


# ProductoUpdateView
class ListadoProductoUpdateView(MaestroUpdateView):
	model = ConfigViews.model
	list_view_name = ConfigViews.list_view_name
	form_class = ConfigViews.form_class
	template_name = ConfigViews.template_form
	success_url = ConfigViews.success_url
	
	#-- Indicar el permiso que requiere para ejecutar la acción.
	permission_required = ConfigViews.permission_change
	
	extra_context = {
		 "accion": f"Editar {ConfigViews.model._meta.verbose_name}",
		"list_view_name" : ConfigViews.list_view_name
	}


# ProductoDeleteView
class ListadoProductoDeleteView (MaestroDeleteView):
	model = ConfigViews.model
	list_view_name = ConfigViews.list_view_name
	template_name = ConfigViews.template_delete
	success_url = ConfigViews.success_url
	
	#-- Indicar el permiso que requiere para ejecutar la acción.
	permission_required = ConfigViews.permission_delete
	
	extra_context = {
		"accion": f"Eliminar {ConfigViews.model._meta.verbose_name}",
		"list_view_name" : ConfigViews.list_view_name,
		"mensaje": "Estás seguro de eliminar el Registro"
	}




class ListadoProductoInformeView(View):
    def get(self, request, *args, **kwargs):
        formato = request.GET.get('formato', 'pdf')
        email = request.GET.get('email', None)
        formatos = request.GET.getlist('formatos')

        if formato == 'pdf':
            return self.generar_pdf()
        elif formato == 'csv':
            return self.generar_csv()
        elif formato == 'word':
            return self.generar_word()
        elif formato == 'excel':
            return self.generar_excel()
        else:
            return HttpResponse("Formato no soportado", status=400)

    def post(self, request, *args, **kwargs):
        email = request.POST.get('email')
        formatos = request.POST.getlist('formatos')

        if email and formatos:
            return self.enviar_email_con_adjuntos(email, formatos)
        
        return JsonResponse({'success': False, 'error': 'Faltan parámetros necesarios.'}, status=400)

    def enviar_email_con_adjuntos(self, email, formatos):
        attachments = []

        if 'pdf' in formatos:
            pdf_response = self.generar_pdf()
            if pdf_response.content:
                attachments.append(('informe_productos.pdf', pdf_response.content, 'application/pdf'))
        
        if 'csv' in formatos:
            csv_response = self.generar_csv()
            if csv_response.content:
                attachments.append(('informe_productos.csv', csv_response.content, 'text/csv'))
        
        if 'word' in formatos:
            word_response = self.generar_word()
            if word_response and word_response.content:
                attachments.append(('informe_productos.docx', word_response.content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        
        if 'excel' in formatos:
            excel_response = self.generar_excel()
            if excel_response.content:
                attachments.append(('informe_productos.xlsx', excel_response.content, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'))

        if attachments:
            try:
                email_message = EmailMessage(
                    subject='Informe de Productos',
                    body='Adjunto encontrará los archivos solicitados en los formatos seleccionados.',
                    from_email=email,  # Cambia esto por tu email
                    to=[email]
                )
                for filename, filedata, mimetype in attachments:
                    email_message.attach(filename, filedata, mimetype)
                email_message.send()
                return JsonResponse({'success': True, 'message': 'Email enviado correctamente.'})
            except Exception as e:
                return JsonResponse({'success': False, 'error': str(e)})
        else:
            return JsonResponse({'success': False, 'error': 'No se generaron adjuntos.'})


    def generar_pdf(self):
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=landscape(letter))
        p.setTitle("Informe de Productos")
        width, height = landscape(letter)

        # Título centrado
        p.setFont("Helvetica-Bold", 16)
        title_text = "Informe de Productos"
        text_width = p.stringWidth(title_text, "Helvetica-Bold", 16)
        p.drawString((width - text_width) / 2, height - 40, title_text)

        # Datos de la tabla
        productos = Producto.objects.all()
        table_headers = ['Código', 'Nombre', 'Tip. Producto', 'Precio', 'Costo', 'Stock']
        data = [table_headers]

        for producto in productos:
            row = [
                producto.codigo_producto,
                producto.nombre_producto,
                producto.tipo_producto,
                producto.precio,
                producto.costo,
                producto.stock
            ]
            data.append(row)

        # Crear y estilizar la tabla
        table = Table(data)
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ]))

        # Centrar la tabla en la página
        table_width, table_height = table.wrap(0, 0)
        x_position = (width - table_width) / 2
        y_position = height - 100 - table_height
        table.drawOn(p, x_position, y_position)

        # Finalizar el PDF
        p.showPage()
        p.save()
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename="informe_productos.pdf"'
        return response
    

    def generar_csv(self):
        buffer = StringIO()
        writer = csv.writer(buffer)
        writer.writerow(['Código', 'Nombre', 'Tipo Producto', 'Precio', 'Costo', 'Stock'])

        productos = Producto.objects.all()
        for producto in productos:
            writer.writerow([
                producto.codigo_producto,
                producto.nombre_producto,
                producto.tipo_producto,
                f"${producto.precio:.2f}",
                f"${producto.costo:.2f}",
                producto.stock
            ])

        # Preparar el contenido del buffer para la respuesta HTTP
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="informe_productos.csv"'
        
        return response

    def generar_word(self):
        buffer = BytesIO()
        doc = Document()
        doc.add_heading('Informe De Productos', level=1)

        # Crear la tabla
        productos = Producto.objects.all()
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Centrar la tabla en el documento

        # Encabezados de la tabla
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Código'
        hdr_cells[1].text = 'Nombre'
        hdr_cells[2].text = 'Tipo'
        hdr_cells[3].text = 'Precio'
        hdr_cells[4].text = 'Costo'
        hdr_cells[5].text = 'Stock'

        # Aplicar borde a la fila de encabezados
        for cell in hdr_cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))

        # Agregar los datos de cada producto en una nueva fila
        for producto in productos:
            row_cells = table.add_row().cells
            row_cells[0].text = str(producto.codigo_producto)
            row_cells[1].text = producto.nombre_producto
            row_cells[2].text = producto.tipo_producto
            row_cells[3].text = f"${producto.precio:.2f}"
            row_cells[4].text = f"${producto.costo:.2f}"
            row_cells[5].text = str(producto.stock)

        # Agregar bordes a toda la tabla
        for row in table.rows:
            for cell in row.cells:
                cell.border_color = '000000'
                tc = cell._element
                tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/><w:left w:val="single" w:sz="12" w:space="0" w:color="auto"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="auto"/><w:right w:val="single" w:sz="12" w:space="0" w:color="auto"/></w:tcBorders>'.format(nsdecls('w'))))

        # Guardar el documento en el buffer
        doc.save(buffer)
        buffer.seek(0)

        # Configurar la respuesta HTTP para descarga del archivo
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="informe_productos.docx"'
        return response
        

    def generar_excel(self):
        buffer = BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.title = "Informe De Productos"

        ws.append(['Código', 'Nombre', 'Tipo Producto', 'Precio', 'Costo', 'Stock'])

        productos = Producto.objects.all()
        for producto in productos:
            ws.append([
                producto.codigo_producto,
                producto.nombre_producto,
                producto.tipo_producto,
                f"${producto.precio:.2f}",
                f"${producto.costo:.2f}",
                producto.stock
            ])

        wb.save(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="informe_productos.xlsx"'
        return response



    